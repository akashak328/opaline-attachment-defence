"""
attachment_converter.py
───────────────────────
Converts potentially malicious email attachments (TXT, DOCX, PDF)
into safe, static PNG images inside a sandboxed process.

Supported formats:
  .txt  → PIL image render
  .docx → page-by-page bitmap via Spire.Doc
  .pdf  → convert to DOCX first, then bitmap render
"""

import os
import textwrap
from typing import Tuple

import docx
from docx import Document
from pdf2docx import parse
from PIL import Image, ImageDraw, ImageFont
from spire.doc import *
from spire.doc.common import *

ATTACHMENTS_DIR = "static/attachments"
FONT_PATH       = "static/arial.ttf"
FONT_SIZE       = 20
WRAP_WIDTH      = 80


# ── TXT → PNG ─────────────────────────────────────────────────────────────────

def text_to_image(text_file: str, fid: int) -> str:
    """
    Renders a plain-text attachment as a PNG image.

    Args:
        text_file : filename of the .txt file inside ATTACHMENTS_DIR
        fid       : unique file/mail ID used to name the output PNG

    Returns:
        output PNG filename (relative to ATTACHMENTS_DIR)
    """
    image_file = f"m{fid}_1.png"

    with open(os.path.join(ATTACHMENTS_DIR, text_file), 'r', encoding='utf-8', errors='replace') as f:
        text = f.read()

    wrapped_text = textwrap.wrap(text, width=WRAP_WIDTH)

    line_height = FONT_SIZE + 5
    img_width   = max((len(line) * FONT_SIZE for line in wrapped_text), default=400) * 2
    img_height  = max(len(wrapped_text) * line_height, 100)

    img  = Image.new("RGB", (img_width, img_height), "white")
    draw = ImageDraw.Draw(img)

    try:
        font = ImageFont.truetype(FONT_PATH, FONT_SIZE)
    except IOError:
        font = ImageFont.load_default()

    y_text = 0
    for line in wrapped_text:
        draw.text((50, y_text), line, font=font, fill="black")
        y_text += line_height

    out_path = os.path.join(ATTACHMENTS_DIR, image_file)
    img.save(out_path)
    print(f"[INFO] TXT → PNG saved: {out_path}")
    return image_file


# ── DOCX → PNG ────────────────────────────────────────────────────────────────

def word_to_img(wfile: str, fid: int) -> int:
    """
    Converts each page of a DOCX file into a PNG image using Spire.Doc.

    Args:
        wfile : filename of the .docx file inside ATTACHMENTS_DIR
        fid   : unique file/mail ID used to name the output PNGs

    Returns:
        number of pages + 1 (counter value after loop)
    """
    document = Document()
    document.LoadFromFile(os.path.join(ATTACHMENTS_DIR, wfile))

    image_streams = document.SaveImageToStreams(ImageType.Bitmap)

    i = 1
    for image in image_streams:
        image_name = f"m{fid}_{i}.png"
        out_path   = os.path.join(ATTACHMENTS_DIR, image_name)
        with open(out_path, 'wb') as image_file:
            image_file.write(image.ToArray())
        print(f"[INFO] DOCX page {i} → PNG saved: {out_path}")
        i += 1

    document.Close()
    return i   # caller uses (i - 1) as page count


# ── PDF → DOCX → PNG ──────────────────────────────────────────────────────────

def convert_pdf_to_docx(input_file: str, output_file: str, pages: Tuple = None) -> dict:
    """
    Converts a PDF to DOCX using pdf2docx.

    Args:
        input_file  : full path to the source PDF
        output_file : full path for the output DOCX
        pages       : optional tuple of page numbers (1-indexed strings)

    Returns:
        summary dict with file info
    """
    page_list = None
    if pages:
        page_list = [int(i) for i in list(pages) if i.isnumeric()]

    result = parse(
        pdf_file=input_file,
        docx_with_path=output_file,
        pages=page_list
    )

    summary = {
        "File":        input_file,
        "Pages":       str(page_list),
        "Output File": output_file,
        "Result":      result
    }
    print("\n".join(f"{k}: {v}" for k, v in summary.items()))
    return summary


def pdf_to_img(pdf_fname: str, fid: int) -> int:
    """
    Full pipeline: PDF → DOCX → PNG images.

    Args:
        pdf_fname : filename of the .pdf file inside ATTACHMENTS_DIR
        fid       : unique file/mail ID

    Returns:
        number of pages converted
    """
    p1        = pdf_fname.split(".")
    docx_name = p1[0] + ".docx"

    pdf_path  = os.path.join(ATTACHMENTS_DIR, pdf_fname)
    docx_path = os.path.join(ATTACHMENTS_DIR, docx_name)

    convert_pdf_to_docx(pdf_path, docx_path)
    page_count = word_to_img(docx_name, fid)
    return page_count - 1


# ── Malicious Content Keyword Scan ────────────────────────────────────────────

def scan_txt_for_keywords(fname: str, testdata_path: str = "static/testdata.txt") -> bool:
    """Check .txt attachment for known malicious keyword patterns."""
    with open(testdata_path, "r") as fp:
        keywords = fp.read().split("|")

    with open(os.path.join(ATTACHMENTS_DIR, fname), "r", encoding="utf-8", errors="replace") as f:
        txt = f.read()

    return any(kw in txt for kw in keywords if kw)


def scan_docx_for_keywords(fname: str, testdata_path: str = "static/testdata.txt") -> bool:
    """Check .docx attachment paragraph content for malicious keywords."""
    with open(testdata_path, "r") as fp:
        keywords = fp.read().split("|")

    doc       = docx.Document(os.path.join(ATTACHMENTS_DIR, fname))
    all_paras = doc.paragraphs

    for kw in keywords:
        for para in all_paras:
            if kw and kw in para.text:
                return True
    return False


def scan_pdf_for_keywords(pdf_fname: str, fid: int, testdata_path: str = "static/testdata.txt") -> bool:
    """Convert PDF → DOCX, then scan paragraphs for malicious keywords."""
    p1        = pdf_fname.split(".")
    docx_name = p1[0] + ".docx"
    pdf_path  = os.path.join(ATTACHMENTS_DIR, pdf_fname)
    docx_path = os.path.join(ATTACHMENTS_DIR, docx_name)

    convert_pdf_to_docx(pdf_path, docx_path)
    return scan_docx_for_keywords(docx_name, testdata_path)
